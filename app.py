import streamlit as st
import sounddevice as sd
import numpy as np
import scipy.io.wavfile as wav
import io
import time
import numpy as np
import Lib.Whisper_extract_text_and_information
import Lib.Extract_top_n_files
import time
import os
import re
import json
from docx import Document

from datetime import datetime

def placeholder_replace_download(filepath,file_name):
    doc = Document(filepath)
    print(filepath)
    print(inputs)

    for key, value in inputs.items():
        for paragraph in doc.paragraphs:
    # Check if the paragraph is a heading (Heading 1, Heading 2, etc.) or contains the value
            if value in paragraph.text or paragraph.style.name.startswith('Heading'):  
                for run in paragraph.runs:
                    if value in run.text:  # Check if the value exists in the run
                        run.text = run.text.replace(value, key, 1)  # Replace only once
                        break  # Exit the inner loop once the replacement is done
                break

    doc.save('final_word/'+file_name)






st.title("Legal template finder using voice")
st.write("")
st.write("")
st.write("")

# Initialize session state for recording
if 'recording' not in st.session_state:
    st.session_state.recording = False
if 'audio_data' not in st.session_state:
    st.session_state.audio_data = []
if 'function_triggered' not in st.session_state:
    st.session_state.function_triggered = False
if 'return_value1' not in st.session_state:
    st.session_state.return_value1 = None 
if 'return_value2' not in st.session_state:
    st.session_state.return_value2 = None
if 'Button_pressed' not in st.session_state:
    st.session_state.Button_pressed = False
if 'Button_pressed1' not in st.session_state:
    st.session_state.Button_pressed1=False 
if 'json_file' not in st.session_state:
    st.session_state.json_file=None

def update_session_state(filename):
    return_value1, return_value2 = Lib.Whisper_extract_text_and_information.extract_information_from_audio(filename)
    
    # Update session state variables
    st.session_state['return_value1'] = return_value1
    st.session_state['return_value2'] = return_value2
    

# Function to record audio
def record_audio():
    st.session_state.audio_data = []  # Reset audio data
    st.info("Recording...")
    
    while st.session_state.recording:
        # Record audio in chunks
        chunk = sd.rec(int(1 * 44100), samplerate=44100, channels=1, dtype='int16')
        sd.wait()  # Wait until the chunk is finished recording
        st.session_state.audio_data.append(chunk)

    st.success("Recording stopped.")


# Create three columns with different widths
col1, col2, col3 = st.columns([1, 8, 1])  # Middle column is wider

# Add buttons to the middle column
with col2:
    # Create a container for buttons
    with st.container():
        # Create three equal-sized columns for the buttons
        button_col1, button_col2, button_col3 = st.columns([1, 1, 1])

        with button_col1:
            if st.button("Start Recording", key="start"):
                if not st.session_state.recording:  # Start recording only if not already recording
                    st.session_state.recording = True
                    record_audio()
                st.write("Recording started!")

        with button_col2:
            if st.button("End Recording", key="end"):
                st.session_state.recording = False
                sd.stop()  # Stop the recording
                if st.session_state.audio_data:

                    suc=st.success("Recording ended!")
                    time.sleep(0.4)
                    suc.empty()
                    # Concatenate recorded chunks into a single NumPy array
                    audio_array = np.concatenate(st.session_state.audio_data, axis=0)  # Ensure it's a NumPy array
                    # Create a BytesIO object for the recorded audio
                    audio_io = io.BytesIO()
                    wav.write(audio_io, 44100, audio_array)  # Write the NumPy array to the BytesIO object
                    audio_io.seek(0)  # Move to the start of the BytesIO buffer
        
                    # Display the audio player
                    
                    current_datetime = datetime.now().strftime('%Y-%m-%d_%H%M%S')
                    filename='Recording'+str(current_datetime)+".mp3"
                    filepath=r'C:\Users\Pc\Downloads'
                    file_path = os.path.join(filepath, filename)
                    # Automatically trigger download of the audio file
                    st.download_button("Download Audio",
                    audio_io,file_name=filename,
                    mime="audio/mp3",key="download_audio",
                    on_click=lambda: st.session_state.update(zip(['return_value1', 'return_value2'],Lib.Whisper_extract_text_and_information.extract_information_from_audio(file_path))))
                    
                else:
                    st.warning("No audio recorded.")

        with button_col3:
            if st.button("Reset", key="reset"):
                st.session_state.recording = False  # Ensure recording is stopped
                suc=st.success("Recording reset.")
                time.sleep(0.4)
                st.session_state.audio_data = []
                st.session_state.return_value=None
                st.session_state.return_value1=None
                st.session_state.Button_pressed=False
                suc.empty()
# Optional: Add some content to the side columns
with col1:
    st.write("")
    
with col3:
    st.write("")


if st.session_state.audio_data:
    st.write("")
    st.write("")
    # Concatenate recorded chunks into a single NumPy array
    audio_array = np.concatenate(st.session_state.audio_data, axis=0)  # Ensure it's a NumPy array
    # Create a BytesIO object for the recorded audio
    audio_io = io.BytesIO()
    wav.write(audio_io, 44100, audio_array)  # Write the NumPy array to the BytesIO object
    audio_io.seek(0)  # Move to the start of the BytesIO buffer
    # Display the audio player
    st.audio(audio_io, format='audio/mp3')



if st.session_state.return_value1:

    st.write("")
    st.markdown("### Returned Value", unsafe_allow_html=True)
    st.write(st.session_state.return_value1)



    st.write("")
    col1, col2, col3 = st.columns([1, 1, 1])
    with col2:
        if st.button("Start Searching", key="search"):
            #st.write(st.session_state.return_value2)
            st.session_state.Button_pressed=True


if st.session_state.Button_pressed==True:
    st.write("")
    filtered_files=Lib.Extract_top_n_files.extract_n_files(st.session_state.return_value1)
    directory = r'word_placeholder/'

    for file in filtered_files:
        file_path = os.path.join(directory, file)

    # Ensure the file exists before creating a download button
    for file in filtered_files:
        if os.path.exists(file_path):
            if st.button(f"Download {file}"):
                print(type(file))
                print(re.sub('.docx','.json',file))
                st.session_state.json_file=re.sub('.docx','.json',file) 
       


if st.session_state.json_file:
    file_name=re.sub('.json','.docx',st.session_state.json_file)
    if st.session_state.Button_pressed1==True:
        if os.path.exists('final_word/'+file_name):
            os.remove(file_path)
            print ('previous file deleted')
            st.session_state.Button_pressed1=False
    with open('json_placeholder/'+st.session_state.json_file,"r") as file:
        data=json.load(file)
        inputs={}
    for key in data.keys():
        inputs[key] = st.text_input(f"Enter {re.sub("<<|>>","",key).capitalize()}", value='')
        
    directory = r'word_placeholder/'
    file_path = os.path.join(directory, file_name)
    if st.button("down", key="down"):
        doc = Document(file_path)
        for key, value in inputs.items():
            for paragraph in doc.paragraphs:
        # Check if the paragraph is a heading (Heading 1, Heading 2, etc.) or contains the value
                if value=='':
                    pass
                elif key in paragraph.text or paragraph.style.name.startswith('Heading'):  
                    for run in paragraph.runs:
                        if key in run.text:  # Check if the value exists in the run
                            run.text = run.text.replace(key,value , 1)
                            print(key,value)  # Replace only once
                            break  # Exit the inner loop once the replacement is done
                    break
        doc.save('final_word/'+file_name)
        st.session_state.Button_pressed1=True
